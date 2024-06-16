import hashlib

import sympy
import random
import math
import docx

from utils import *


class RSA_CRIPT:
    @staticmethod
    def generate_large_prime(bit_length):
        while True:
            prime_candidate = sympy.randprime(2 ** (bit_length - 1), 2 ** bit_length)
            if sympy.isprime(prime_candidate):
                return prime_candidate

    @staticmethod
    def euclid(a, n):
        g_arr = []
        u_arr = []
        v_arr = []
        g_arr.append(n)
        g_arr.append(a)
        u_arr.append(1)
        u_arr.append(0)
        v_arr.append(0)
        v_arr.append(1)
        while g_arr[len(g_arr) - 1] != 1:
            len_ = len(g_arr) - 1
            y = g_arr[len_ - 1] // g_arr[len_]
            g_arr.append(g_arr[len_ - 1] - y * g_arr[len_])
            u_arr.append(u_arr[len_ - 1] - y * u_arr[len_])
            v_arr.append(v_arr[len_ - 1] - y * v_arr[len_])
        if v_arr[len(v_arr) - 1] < 0:
            return n + v_arr[len(v_arr) - 1]
        else:
            return v_arr[len(v_arr) - 1]

    def key(self, bit):
        p = self.generate_large_prime(bit)
        q = self.generate_large_prime(bit)
        n = q * p
        f = (p - 1) * (q - 1)
        arr_key = [n]
        e = self.generate_large_prime(random.randint(bit - 200, bit + 200))
        arr_key.append(e)
        d = self.euclid(e, f)
        arr_key.append(d)
        return arr_key

    def eport_key(self, bit, link_key_public, link_key_private):
        arr_key = self.key(bit // 2)
        link_key_public = set_link(link_key_public).split('.')
        link_key_private = set_link(link_key_private).split('.')
        key_public = str(arr_key[1]) + " " + str(arr_key[0])
        key_private = str(arr_key[2]) + " " + str(arr_key[0])
        self.save_file(link_key_public[0], '.pem', key_public)
        self.save_file(link_key_private[0], '.pem', key_private)
        keyname_public = get_keyname(link_key_public[0], str(bit)) + "~"
        keyname_private = get_keyname(link_key_private[0], str(bit)) + "~"
        insert_data("\n" + "Public~" + keyname_public + key_public)
        insert_data("\n" + "Private~" + keyname_private + key_private)

    @staticmethod
    def int_to_n93(n):
        n93 = ""
        while n != 0:
            nf = n % 93
            n93 = chr(nf + 33) + n93
            n //= 93
        return n93

    @staticmethod
    def n93_to_int(n93):
        num_int = 0
        p = 0
        for i in range(-1, -len(n93) - 1, -1):
            num_int += (ord(n93[i]) - 33) * pow(93, p)
            p += 1
        return num_int

    @staticmethod
    def powmod(a, b, m):
        res = 1
        while b:
            if b % 2 == 1:
                res *= a
                res %= m
            a *= a
            a %= m
            b //= 2
        return res

    @staticmethod
    def full_word(w):
        w = hex(ord(w))
        w = w[2:]
        while len(w) < 4:
            w = '0' + w
        return w

    def news_block(self, str_ban_ro, bit):
        ln = bit // 4 - 5
        str_block = ""
        for i in str_ban_ro:
            str_block += self.full_word(i)
        num_block = math.ceil(len(str_block) / ln)
        for i in range(ln - len(str_block) % ln):
            str_block += '0'
        arr_block = []
        for i in range(num_block):
            c_block = "ffff" + str_block[i * ln: (i + 1) * ln]
            arr_block.append(int(c_block, 16))
        return arr_block

    def encrypt(self, str_ban_ro, str_key_public):
        key_public = str_key_public.split()
        str_ban_ma = ""
        arr_block = self.news_block(str_ban_ro, round(math.log2(int(key_public[1]))))
        len_ = len(arr_block)
        # per = 0
        for i in range(len_):
            str_ban_ma += self.int_to_n93(self.powmod(arr_block[i], int(key_public[0]), int(key_public[1]))) + "~"
            """per_ = round(((i + 1) / len_) * 100)
            if per_ % 10 == 0 and per != per_:
                print(str(per_) + '%', end=' ')
                per = per_"""
        return str_ban_ma

    def Cripting(self, str_ban_ro, link_save, key_public):
        round(math.log2(int(key_public[1])))
        str_ban_ma = self.encrypt(str_ban_ro, key_public)

        hash_w = hashlib.sha256(str_ban_ma.encode("utf-8")).hexdigest()

        hash_w_cript = self.encrypt(hash_w, key_public)
        self.save_file(link_save, '.rsa', str_ban_ma + str(chr(94 + 33)) + hash_w_cript)

    def RSA_Ma_Hoa(self, link_key_public, link_send, link_save):
        link_key_public = set_link(link_key_public)
        link_send = set_link(link_send)
        link_save = set_link(link_save).split('.')
        str_ban_ro = self.input_text(link_send)
        key_public = self.open_file(link_key_public)
        self.Cripting(str_ban_ro, link_save[0], key_public)
        keyname_public = get_keyname(link_key_public[0], "None") + "~"
        insert_data("\n" + "Public~" + keyname_public + key_public)

    def RSA_Ma_Hoa_key(self, key_public, link_send, link_save):
        link_send = set_link(link_send)
        link_save = set_link(link_save).split('.')
        str_ban_ro = self.input_text(link_send)
        self.Cripting(str_ban_ro, link_save[0], key_public)

    @staticmethod
    def input_text(link_dox):
        doc = docx.Document(link_dox)

        # Lặp qua các đoạn văn bản và hiển thị nội dung
        str_input = ""
        for paragraph in doc.paragraphs:
            str_input += paragraph.text + "\n"
        return str_input

    @staticmethod
    def correct_str(cr_str):
        id = 0
        str_res = ""
        while True:
            if id == len(cr_str):
                break
            if cr_str[id] == 'f':
                test_ = 0
                for i in range(4):
                    if cr_str[i + id] != 'f':
                        test_ = 1
                        break
                if test_ == 0:
                    id += 4
                    continue
                else:
                    str_res += cr_str[id]
                    id += 1
                    continue
            else:
                str_res += cr_str[id]
                id += 1
        return str_res

    @staticmethod
    def save_giai_ma(str_giai_ma, link_save):
        doc = docx.Document()
        paragraphs = str_giai_ma.split('\n')
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
        doc.save(link_save)

    def encrypt_(self, str_ban_ma, str_key_private):
        key_private = str_key_private.split()
        str_ban_ma = str_ban_ma.split('~')
        news_block = ""
        str_giai_ma = ""
        len_ = len(str_ban_ma)
        per = 0
        for i in range(len_):
            if str_ban_ma[i] != "":
                c_block = hex(self.powmod(self.n93_to_int(str_ban_ma[i]), int(key_private[0]), int(key_private[1])))
                news_block += c_block[2:]
                per_ = round(((i + 1) / len_) * 100)
                if per_ % 10 == 0 and per != per_:
                    print(str(per_) + '%', end=' ')
                    per = per_
        cr_str = self.correct_str(news_block)
        for i in range(len(cr_str) // 4):
            div_block = cr_str[i * 4: (i + 1) * 4]
            if div_block == "0000":
                break
            str_giai_ma += chr(int(div_block, 16))
        return str_giai_ma

    def decrypt(self, str_ban_ma, str_key_private):
        key_private = str_key_private.split()
        str_ban_ma = str_ban_ma.split('~')
        news_block = ""
        str_giai_ma = ""
        len_ = len(str_ban_ma)
        for i in range(len_):
            if str_ban_ma[i] != "":
                c_block = hex(self.powmod(self.n93_to_int(str_ban_ma[i]), int(key_private[0]), int(key_private[1])))
                news_block += c_block[2:]
        for i in range(len(news_block) // 4):
            div_block = news_block[i * 4: (i + 1) * 4]
            if div_block == '0000':
                break
            try:
                str_giai_ma += chr(int(div_block, 16))
            except ValueError:
                return "-1"
        return str_giai_ma

    def slove_cript(self, link_slove, link_save, key_private):

        # Lặp qua các đoạn văn bản và hiển thị nội dung
        str_word = self.open_file(link_slove)
        str_ban_ma = ""
        hash_w_cript = ""
        bo = 0
        for i in str_word:
            if i == chr(94 + 33):
                bo = 1
                continue
            if bo == 0:
                str_ban_ma += i
            else:
                hash_w_cript += i
        self.decrypt(hash_w_cript, key_private)
        hash_w = hashlib.sha256(str_ban_ma.encode("utf-8")).hexdigest()
        hash_w_test = self.decrypt(hash_w_cript, key_private)
        if hash_w == hash_w_test[1:]:
            str_giai_ma = self.encrypt_(str_ban_ma, key_private)
            if str_giai_ma != "-1":
                self.save_giai_ma(str_giai_ma, link_save)
                return 0
            else:
                return 1
        else:
            return 2

    def RSA_Giai_Ma(self, link_giai_ma, link_save, link_key_private):
        link_giai_ma = set_link(link_giai_ma)
        link_save = set_link(link_save).split('.')
        link_save = link_save[0] + ".docx"
        link_key_private = set_link(link_key_private)
        key_private = self.open_file(link_key_private)
        insert_data("\nPrivate" + get_keyname(link_key_private, "None") + key_private)
        return self.slove_cript(link_giai_ma, link_save, key_private)

    def RSA_Giai_Ma_Key(self, link_giai_ma, link_save, key_private):
        link_giai_ma = set_link(link_giai_ma)
        link_save = set_link(link_save).split('.')
        link_save = link_save[0] + ".docx"
        return self.slove_cript(link_giai_ma, link_save, key_private)

    @staticmethod
    def save_file(ten_file, duoi, string):
        with open(ten_file + duoi, 'w', encoding='utf-8') as file_:
            file_.write(string)

    @staticmethod
    def open_file(ten_file):
        with open(ten_file, 'r', encoding='utf-8') as file_:
            return file_.read()
