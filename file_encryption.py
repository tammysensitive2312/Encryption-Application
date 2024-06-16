import csv
import cv2
import numpy as np
from Cryptodome.Cipher import AES, PKCS1_OAEP
from Cryptodome.PublicKey import RSA
from Cryptodome.Random import get_random_bytes
import wave
import docx


class FileEncryptor:
    def __init__(self, rsa_public_key, rsa_private_key):
        self.rsa_public_key = RSA.import_key(rsa_public_key)
        self.rsa_private_key = RSA.import_key(rsa_private_key)

    def encrypt(self, input_file_path, output_file_path):
        raise NotImplementedError("This method needs to be overridden.")

    def decrypt(self, input_file_path, output_file_path):
        raise NotImplementedError("This method needs to be overridden.")

    def rsa_encrypt(self, aes_key):
        cipher_rsa = PKCS1_OAEP.new(self.rsa_public_key)
        return cipher_rsa.encrypt(aes_key)

    def rsa_decrypt(self, encrypted_aes_key):
        cipher_rsa = PKCS1_OAEP.new(self.rsa_private_key)
        return cipher_rsa.decrypt(encrypted_aes_key)


class PDFEncryptor(FileEncryptor):
    def encrypt(self, input_file_path, output_file_path):
        aes_key = get_random_bytes(32)
        aes_iv = get_random_bytes(16)

        with open(input_file_path, 'rb') as file:
            plaintext = file.read()

        cipher = AES.new(aes_key, AES.MODE_CFB, iv=aes_iv)
        ciphertext = cipher.encrypt(plaintext)

        encrypted_aes_key = self.rsa_encrypt(aes_key)

        with open(output_file_path, 'wb') as file:
            file.write(encrypted_aes_key)
            file.write(aes_iv)
            file.write(ciphertext)

    def decrypt(self, input_file_path, output_file_path):
        with open(input_file_path, 'rb') as file:
            encrypted_aes_key = file.read(256)
            aes_iv = file.read(16)
            ciphertext = file.read()

        aes_key = self.rsa_decrypt(encrypted_aes_key)
        cipher = AES.new(aes_key, AES.MODE_CFB, iv=aes_iv)
        plaintext = cipher.decrypt(ciphertext)

        with open(output_file_path, 'wb') as file:
            file.write(plaintext)


class WavEncryptor(FileEncryptor):
    def pad(self, data):
        block_size = AES.block_size
        return data + (block_size - len(data) % block_size) * b"\0"

    def unpad(self, data):
        return data.rstrip(b"\0")

    def encrypt(self, input_file_path, output_file_path):
        aes_key = get_random_bytes(32)
        cipher_aes = AES.new(aes_key, AES.MODE_ECB)

        with wave.open(input_file_path, 'rb') as original_wave:
            n_channels = original_wave.getnchannels()
            sampwidth = original_wave.getsampwidth()
            framerate = original_wave.getframerate()
            n_frames = original_wave.getnframes()
            audio_data = original_wave.readframes(n_frames)

        padded_audio = self.pad(audio_data)
        obfuscated_audio = cipher_aes.encrypt(padded_audio)

        encrypted_aes_key = self.rsa_encrypt(aes_key)

        with wave.open(output_file_path, 'wb') as output_wave:
            output_wave.setnchannels(n_channels)
            output_wave.setsampwidth(sampwidth)
            output_wave.setframerate(framerate)
            output_wave.writeframes(encrypted_aes_key + obfuscated_audio)

    def decrypt(self, input_file_path, output_file_path):
        with wave.open(input_file_path, 'rb') as obfuscated_wave:
            encrypted_aes_key = obfuscated_wave.readframes(256 // obfuscated_wave.getsampwidth())
            obfuscated_audio = obfuscated_wave.readframes(-1)
            sample_width = obfuscated_wave.getsampwidth()
            num_channels = obfuscated_wave.getnchannels()
            framerate = obfuscated_wave.getframerate()

        aes_key = self.rsa_decrypt(encrypted_aes_key)
        cipher_aes = AES.new(aes_key, AES.MODE_ECB)
        padded_audio = cipher_aes.decrypt(obfuscated_audio)
        original_audio = self.unpad(padded_audio)

        with wave.open(output_file_path, 'wb') as output_wave:
            output_wave.setnchannels(num_channels)
            output_wave.setsampwidth(sample_width)
            output_wave.setframerate(framerate)
            output_wave.writeframes(original_audio)


class ImageEncryptor(FileEncryptor):
    def pad(self, data):
        block_size = AES.block_size
        padding_len = block_size - len(data) % block_size
        return data + padding_len * chr(padding_len).encode()

    def unpad(self, data):
        padding_len = data[-1]
        return data[:-padding_len]

    def encrypt(self, input_file_path, output_file_path):
        aes_key = get_random_bytes(32)
        cipher_aes = AES.new(aes_key, AES.MODE_ECB)

        image = cv2.imread(input_file_path)
        h, w, c = image.shape
        image_data = image.tobytes()

        padded_image_data = self.pad(image_data)
        encrypted_image_data = cipher_aes.encrypt(padded_image_data)

        encrypted_aes_key = self.rsa_encrypt(aes_key)

        with open(output_file_path, 'wb') as file:
            file.write(encrypted_aes_key)
            file.write(h.to_bytes(4, 'big'))
            file.write(w.to_bytes(4, 'big'))
            file.write(c.to_bytes(4, 'big'))
            file.write(encrypted_image_data)

    def decrypt(self, input_file_path, output_file_path):
        with open(input_file_path, 'rb') as file:
            encrypted_aes_key = file.read(256)
            h = int.from_bytes(file.read(4), 'big')
            w = int.from_bytes(file.read(4), 'big')
            c = int.from_bytes(file.read(4), 'big')
            encrypted_image_data = file.read()

        aes_key = self.rsa_decrypt(encrypted_aes_key)
        cipher_aes = AES.new(aes_key, AES.MODE_ECB)
        padded_image_data = cipher_aes.decrypt(encrypted_image_data)
        image_data = self.unpad(padded_image_data)

        image = np.frombuffer(image_data, dtype=np.uint8).reshape((h, w, c))
        cv2.imwrite(output_file_path, image)


class CsvEncryptor(FileEncryptor):
    def pad(self, data):
        block_size = AES.block_size
        padding_len = block_size - len(data) % block_size
        return data + padding_len * chr(padding_len).encode()

    def unpad(self, data):
        padding_len = data[-1]
        return data[:-padding_len]

    def encrypt(self, input_file_path, output_file_path):
        aes_key = get_random_bytes(32)
        cipher_aes = AES.new(aes_key, AES.MODE_ECB)

        with open(input_file_path, 'r', newline='') as file:
            reader = csv.reader(file)
            csv_data = "\n".join([",".join(row) for row in reader]).encode()

        padded_csv_data = self.pad(csv_data)
        encrypted_csv_data = cipher_aes.encrypt(padded_csv_data)

        encrypted_aes_key = self.rsa_encrypt(aes_key)

        with open(output_file_path, 'wb') as file:
            file.write(encrypted_aes_key)
            file.write(encrypted_csv_data)

    def decrypt(self, input_file_path, output_file_path):
        with open(input_file_path, 'rb') as file:
            encrypted_aes_key = file.read(256)
            encrypted_csv_data = file.read()

        aes_key = self.rsa_decrypt(encrypted_aes_key)
        cipher_aes = AES.new(aes_key, AES.MODE_ECB)
        padded_csv_data = cipher_aes.decrypt(encrypted_csv_data)
        csv_data = self.unpad(padded_csv_data).decode()

        with open(output_file_path, 'w', newline='') as file:
            writer = csv.writer(file)
            for row in csv_data.split("\n"):
                writer.writerow(row.split(","))


class WordEncryptor(FileEncryptor):
    def pad(self, data):
        block_size = AES.block_size
        padding_len = block_size - len(data) % block_size
        return data + padding_len * chr(padding_len).encode()

    def unpad(self, data):
        padding_len = data[-1]
        return data[:-padding_len]

    def encrypt(self, input_file_path, output_file_path):
        aes_key = get_random_bytes(32)
        cipher_aes = AES.new(aes_key, AES.MODE_ECB)

        document = docx.Document(input_file_path)
        docx_data = "\n".join([para.text for para in document.paragraphs]).encode()

        padded_docx_data = self.pad(docx_data)
        encrypted_docx_data = cipher_aes.encrypt(padded_docx_data)

        encrypted_aes_key = self.rsa_encrypt(aes_key)

        with open(output_file_path, 'wb') as file:
            file.write(encrypted_aes_key)
            file.write(encrypted_docx_data)

    def decrypt(self, input_file_path, output_file_path):
        with open(input_file_path, 'rb') as file:
            encrypted_aes_key = file.read(256)
            encrypted_docx_data = file.read()

        aes_key = self.rsa_decrypt(encrypted_aes_key)
        cipher_aes = AES.new(aes_key, AES.MODE_ECB)
        padded_docx_data = cipher_aes.decrypt(encrypted_docx_data)
        docx_data = self.unpad(padded_docx_data).decode()

        document = docx.Document()
        for line in docx_data.split("\n"):
            document.add_paragraph(line)
        document.save(output_file_path)


class EncryptorFactory:
    @staticmethod
    def get_encryptor(file_type, rsa_public_key, rsa_private_key):
        if file_type == "pdf":
            return PDFEncryptor(rsa_public_key, rsa_private_key)
        elif file_type == "wav":
            return WavEncryptor(rsa_public_key, rsa_private_key)
        elif file_type == "csv":
            return CsvEncryptor(rsa_public_key, rsa_private_key)
        elif file_type == "png" or file_type == "jpg" or file_type == "jpeg":
            return ImageEncryptor(rsa_public_key, rsa_private_key)
        elif file_type == "docx":
            return WordEncryptor(rsa_public_key, rsa_private_key)
        else:
            raise ValueError("Unsupported file type")
